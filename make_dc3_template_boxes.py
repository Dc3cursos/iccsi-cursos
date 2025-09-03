from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def _add_heading(doc: Document, text: str, size: int = 14) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_text(doc: Document, text: str, size: int = 11, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def _add_box_row(doc: Document, label: str, key_prefix: str, total_cells: int) -> None:
    # Etiqueta
    _add_text(doc, label, bold=True)
    # Tabla de una fila con total_cells columnas
    table = doc.add_table(rows=1, cols=total_cells)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i in range(total_cells):
        cell = table.cell(0, i)
        # Borde de celda
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '8')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), '000000')
            tcBorders.append(el)
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{{{{ {key_prefix}_{i+1} }}}}")
        run.font.size = Pt(11)


def build_boxes_template(path: Path) -> None:
    doc = Document()
    _add_heading(doc, 'FORMATO DC-3')
    _add_heading(doc, 'CONSTANCIA DE COMPETENCIAS O DE HABILIDADES LABORALES')

    _add_heading(doc, 'DATOS DEL TRABAJADOR')
    _add_text(doc, 'Nombre (Apellido paterno, Apellido materno, Nombre[s]): {{ APELLIDO_PATERNO }} {{ APELLIDO_MATERNO }} {{ NOMBRES }}')
    _add_text(doc, 'Puesto: {{ PUESTO }}')
    _add_text(doc, 'Ocupación específica (Catálogo Nacional de Ocupaciones): {{ OCUPACION }}')
    _add_text(doc, 'CURP:')
    _add_box_row(doc, 'CURP', 'CURP', 18)

    _add_heading(doc, 'DATOS DE LA EMPRESA')
    _add_text(doc, 'Nombre o razón social: {{ RAZON_SOCIAL }}')
    _add_text(doc, 'RFC:')
    _add_box_row(doc, 'RFC', 'RFC', 13)

    _add_heading(doc, 'DATOS DEL PROGRAMA DE CAPACITACIÓN, ADIESTRAMIENTO Y PRODUCTIVIDAD')
    _add_text(doc, 'Nombre del curso: {{ NOMBRE_CURSO }}')
    _add_text(doc, 'Duración en horas: {{ HORAS }}')

    _add_text(doc, 'Periodo de ejecución - INICIO (DÍA/MES/AÑO):')
    _add_box_row(doc, 'DÍA', 'DIA_INICIO', 2)
    _add_box_row(doc, 'MES', 'MES_INICIO', 2)
    _add_box_row(doc, 'AÑO', 'ANIO_INICIO', 4)

    _add_text(doc, 'Periodo de ejecución - FIN (DÍA/MES/AÑO):')
    _add_box_row(doc, 'DÍA', 'DIA_FIN', 2)
    _add_box_row(doc, 'MES', 'MES_FIN', 2)
    _add_box_row(doc, 'AÑO', 'ANIO_FIN', 4)

    _add_text(doc, 'Nombre del agente capacitador o STPS: {{ AGENTE_CAPACITADOR }}   Registro: {{ REGISTRO_AGENTE }}')

    _add_heading(doc, 'FIRMAS')
    _add_text(doc, 'Instructor o tutor: {{ INSTRUCTOR_NOMBRE }}  |  Nombre y firma')
    _add_text(doc, 'Patrón o representante legal: {{ EMPRESA_REPRESENTANTE_LEGAL }}  |  Nombre y firma')
    _add_text(doc, 'Representante de los trabajadores: {{ EMPRESA_REPRESENTANTE_TRABAJADORES }}  |  Nombre y firma')

    _add_heading(doc, 'OTROS')
    _add_text(doc, 'Organización: {{ ORGANIZACION_NOMBRE }}')
    _add_text(doc, 'Fecha de emisión: {{ FECHA_EMISION }}  |  Vigencia: {{ VIGENCIA }}')
    _add_text(doc, 'LOGO: {{ LOGO }}')

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path.as_posix())


if __name__ == '__main__':
    out = Path('plantillas') / 'dc3_plantilla_recuadros.docx'
    build_boxes_template(out)
    print(f'Wrote: {out.as_posix()}')


