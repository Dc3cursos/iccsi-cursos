import json
import re
import sys
from pathlib import Path

from docx import Document


TAG_PATTERN = re.compile(r"\{\{[\s\S]*?\}\}")


def extract_text_blocks(doc: Document):
    # Paragraphs
    for p in doc.paragraphs:
        text = p.text or ""
        if text:
            yield ("paragraph", text)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text or ""
                if text:
                    yield ("cell", text)


def main(path: Path):
    doc = Document(path.as_posix())
    tags = []
    unique = {}
    for kind, text in extract_text_blocks(doc):
        for m in TAG_PATTERN.finditer(text):
            tag = m.group(0)
            snippet_start = max(0, m.start() - 30)
            snippet_end = min(len(text), m.end() + 30)
            snippet = text[snippet_start:snippet_end]
            tags.append({
                "kind": kind,
                "tag": tag,
                "context": snippet,
            })
            key = tag
            unique.setdefault(key, 0)
            unique[key] += 1

    print(json.dumps({
        "file": path.as_posix(),
        "total_tags": len(tags),
        "unique_tags": len(unique),
        "by_tag": sorted(
            [{"tag": k, "count": v} for k, v in unique.items()],
            key=lambda x: x["tag"],
        ),
        "examples": tags[:50],
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python tools/list_docx_raw_tags.py <file.docx>")
        sys.exit(2)
    target = Path(sys.argv[1])
    main(target)


