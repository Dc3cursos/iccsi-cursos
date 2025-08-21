#!/usr/bin/env python3
"""
Script para reemplazar automáticamente variables inválidas en templates Word
con variables válidas para docxtpl.
"""

import sys
import os
from pathlib import Path
from docx import Document
import re

def fix_template_variables(input_file, output_file):
    """
    Reemplaza variables inválidas en un template Word con variables válidas.
    """
    print(f"Procesando: {input_file}")
    
    # Cargar el documento
    doc = Document(input_file)
    
    # Mapeo de reemplazos
    replacements = {
        # Información del trabajador
        r'\{\{Nombre Anotar apellido paterno, apellido materno y nombres\}\}': '{{ APELLIDO_PATERNO }} {{ APELLIDO_MATERNO }} {{ NOMBRES }}',
        r'\{\{trabajador_nombre\}\}': '{{ APELLIDO_PATERNO }} {{ APELLIDO_MATERNO }} {{ NOMBRES }}',
        
        # Información del curso
        r'\{\{Nombre del curso\}\}': '{{ NOMBRE_CURSO }}',
        r'\{\{Duración en horas\}\}': '{{ HORAS }}',
        
        # Información de la empresa
        r'\{\{Nombre o razón social \(En caso de persona física, anotar apellido paterno, apellido materno y nombres\}\}': '{{ RAZON_SOCIAL }}',
        r'\{\{empresa_razon_social\}\}': '{{ RAZON_SOCIAL }}',
        r'\{\{Registro Federal de Contribuyentes con homoclave SHCP\}\}': '{{ RFC }}',
        
        # Información personal
        r'\{\{Clave Única de Registro de Población\}\}': '{{ CURP }}',
        r'\{\{Puesto\}\}': '{{ PUESTO }}',
        r'\{\{ocupacion\}\}': '{{ OCUPACION }}',
        
        # Fechas
        r'\{\{dia_inicio\}\}': '{{ DIA_INICIO }}',
        r'\{\{mes_inicio\}\}': '{{ MES_INICIO }}',
        r'\{\{anio_inicio\}\}': '{{ ANIO_INICIO }}',
        r'\{\{dia_fin\}\}': '{{ DIA_FIN }}',
        r'\{\{mes_fin\}\}': '{{ MES_FIN }}',
        r'\{\{anio_fin\}\}': '{{ ANIO_FIN }}',
        
        # Dígitos individuales CURP
        r'\{\{curp_chars\[0\]\}\}': '{{ CURP_1 }}',
        r'\{\{curp_chars\[1\]\}\}': '{{ CURP_2 }}',
        r'\{\{curp_chars\[2\]\}\}': '{{ CURP_3 }}',
        r'\{\{curp_chars\[3\]\}\}': '{{ CURP_4 }}',
        r'\{\{curp_chars\[4\]\}\}': '{{ CURP_5 }}',
        r'\{\{curp_chars\[5\]\}\}': '{{ CURP_6 }}',
        r'\{\{curp_chars\[6\]\}\}': '{{ CURP_7 }}',
        r'\{\{curp_chars\[7\]\}\}': '{{ CURP_8 }}',
        r'\{\{curp_chars\[8\]\}\}': '{{ CURP_9 }}',
        r'\{\{curp_chars\[9\]\}\}': '{{ CURP_10 }}',
        r'\{\{curp_chars\[10\]\}\}': '{{ CURP_11 }}',
        r'\{\{curp_chars\[11\]\}\}': '{{ CURP_12 }}',
        r'\{\{curp_chars\[12\]\}\}': '{{ CURP_13 }}',
        r'\{\{curp_chars\[13\]\}\}': '{{ CURP_14 }}',
        r'\{\{curp_chars\[14\]\}\}': '{{ CURP_15 }}',
        r'\{\{curp_chars\[15\]\}\}': '{{ CURP_16 }}',
        r'\{\{curp_chars\[16\]\}\}': '{{ CURP_17 }}',
        r'\{\{curp_chars\[17\]\}\}': '{{ CURP_18 }}',
        
        # Dígitos individuales RFC
        r'\{\{rfc_chars\[0\]\}\}': '{{ RFC_1 }}',
        r'\{\{rfc_chars\[1\]\}\}': '{{ RFC_2 }}',
        r'\{\{rfc_chars\[2\]\}\}': '{{ RFC_3 }}',
        r'\{\{rfc_chars\[3\]\}\}': '{{ RFC_4 }}',
        r'\{\{rfc_chars\[4\]\}\}': '{{ RFC_5 }}',
        r'\{\{rfc_chars\[5\]\}\}': '{{ RFC_6 }}',
        r'\{\{rfc_chars\[6\]\}\}': '{{ RFC_7 }}',
        r'\{\{rfc_chars\[7\]\}\}': '{{ RFC_8 }}',
        r'\{\{rfc_chars\[8\]\}\}': '{{ RFC_9 }}',
        r'\{\{rfc_chars\[9\]\}\}': '{{ RFC_10 }}',
        r'\{\{rfc_chars\[10\]\}\}': '{{ RFC_11 }}',
        r'\{\{rfc_chars\[11\]\}\}': '{{ RFC_12 }}',
        r'\{\{rfc_chars\[12\]\}\}': '{{ RFC_13 }}',
        
        # Dígitos de fechas
        r'\{\{dia_inicio_d\[1\]\}\}': '{{ DIA_INICIO_1 }}',
        r'\{\{dia_inicio_d\[2\]\}\}': '{{ DIA_INICIO_2 }}',
        r'\{\{mes_inicio_d\[1\]\}\}': '{{ MES_INICIO_1 }}',
        r'\{\{mes_inicio_d\[2\]\}\}': '{{ MES_INICIO_2 }}',
        r'\{\{anio_inicio_d\[1\]\}\}': '{{ ANIO_INICIO_1 }}',
        r'\{\{anio_inicio_d\[2\]\}\}': '{{ ANIO_INICIO_2 }}',
        r'\{\{anio_inicio_d\[3\]\}\}': '{{ ANIO_INICIO_3 }}',
        r'\{\{anio_inicio_d\[4\]\}\}': '{{ ANIO_INICIO_4 }}',
        r'\{\{dia_fin_d\[1\]\}\}': '{{ DIA_FIN_1 }}',
        r'\{\{dia_fin_d\[2\]\}\}': '{{ DIA_FIN_2 }}',
        r'\{\{mes_fin_d\[1\]\}\}': '{{ MES_FIN_1 }}',
        r'\{\{mes_fin_d\[2\]\}\}': '{{ MES_FIN_2 }}',
        r'\{\{anio_fin_d\[1\]\}\}': '{{ ANIO_FIN_1 }}',
        r'\{\{anio_fin_d\[3\]\}\}': '{{ ANIO_FIN_3 }}',
        r'\{\{anio_fin_d\[4\]\}\}': '{{ ANIO_FIN_4 }}',
        
        # Otros
        r'\{\{5\}\}': '{{ HORAS }}',
        r'\{\{Nombre del agente capacitador o STPS 3/ /     / Fraternidad Migratoria A\.C\. Registro: Fmi111006-4q2-0013\}\}': '{{ AGENTE_CAPACITADOR }}',
    }
    
    # Contador de reemplazos
    total_replacements = 0
    
    # Procesar párrafos
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        for pattern, replacement in replacements.items():
            if re.search(pattern, original_text):
                new_text = re.sub(pattern, replacement, original_text)
                if new_text != original_text:
                    paragraph.text = new_text
                    total_replacements += 1
                    print(f"  Reemplazado en párrafo: {pattern[:50]}...")
    
    # Procesar tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    for pattern, replacement in replacements.items():
                        if re.search(pattern, original_text):
                            new_text = re.sub(pattern, replacement, original_text)
                            if new_text != original_text:
                                paragraph.text = new_text
                                total_replacements += 1
                                print(f"  Reemplazado en tabla: {pattern[:50]}...")
    
    # Procesar headers y footers
    for section in doc.sections:
        # Headers
        if section.header:
            for paragraph in section.header.paragraphs:
                original_text = paragraph.text
                for pattern, replacement in replacements.items():
                    if re.search(pattern, original_text):
                        new_text = re.sub(pattern, replacement, original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
                            total_replacements += 1
                            print(f"  Reemplazado en header: {pattern[:50]}...")
        
        # Footers
        if section.footer:
            for paragraph in section.footer.paragraphs:
                original_text = paragraph.text
                for pattern, replacement in replacements.items():
                    if re.search(pattern, original_text):
                        new_text = re.sub(pattern, replacement, original_text)
                        if new_text != original_text:
                            paragraph.text = new_text
                            total_replacements += 1
                            print(f"  Reemplazado en footer: {pattern[:50]}...")
    
    # Guardar el documento
    doc.save(output_file)
    print(f"Archivo guardado como: {output_file}")
    print(f"Total de reemplazos realizados: {total_replacements}")
    
    return total_replacements

def main():
    if len(sys.argv) != 3:
        print("Uso: python fix_template_variables.py <archivo_entrada> <archivo_salida>")
        print("Ejemplo: python fix_template_variables.py plantillas/CERTIFICADO_DC3_corregido_reparada.docx plantillas/CERTIFICADO_DC3_final.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"Error: El archivo {input_file} no existe.")
        sys.exit(1)
    
    try:
        replacements = fix_template_variables(input_file, output_file)
        print(f"\n✅ Proceso completado exitosamente!")
        print(f"Se realizaron {replacements} reemplazos.")
        print(f"Archivo final: {output_file}")
    except Exception as e:
        print(f"Error durante el proceso: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
