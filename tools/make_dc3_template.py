from docx import Document
from docx.shared import Pt
from pathlib import Path


CONTENT = [
    (True, 'FORMATO DC-3'),
    (True, 'CONSTANCIA DE COMPETENCIAS O DE HABILIDADES LABORALES'),
    (True, 'DATOS DEL TRABAJADOR'),
    (False, 'Nombre (Apellido paterno, Apellido materno, Nombre(s)): {{ APELLIDO_PATERNO }} {{ APELLIDO_MATERNO }} {{ NOMBRES }}'),
    (False, 'Clave Única de Registro de Población (CURP): {{ CURP }}'),
    (False, 'Ocupación específica (Catálogo Nacional de Ocupaciones): {{ OCUPACION }}'),
    (False, 'Puesto: {{ PUESTO }}'),
    (True, 'DATOS DE LA EMPRESA'),
    (False, 'Nombre o razón social: {{ RAZON_SOCIAL }}'),
    (False, 'Registro Federal de Contribuyentes con homoclave (SHCP): {{ RFC }}'),
    (True, 'DATOS DEL PROGRAMA DE CAPACITACIÓN, ADIESTRAMIENTO Y PRODUCTIVIDAD'),
    (False, 'Nombre del curso: {{ NOMBRE_CURSO }}'),
    (False, 'Duración en horas: {{ HORAS }}'),
    (False, 'Periodo de ejecución: Inicio: {{ DIA_INICIO }}/{{ MES_INICIO }}/{{ ANIO_INICIO }}  Fin: {{ DIA_FIN }}/{{ MES_FIN }}/{{ ANIO_FIN }}'),
    (False, 'Área temática del curso: {{ AREA_TEMATICA }}'),
    (False, 'Nombre del agente capacitador o STPS: {{ AGENTE_CAPACITADOR }}  Registro: {{ REGISTRO_AGENTE }}'),
    (True, 'FIRMAS'),
    (False, 'Instructor o tutor: {{ INSTRUCTOR_NOMBRE }}  |  Nombre y firma'),
    (False, 'Patrón o representante legal: {{ EMPRESA_REPRESENTANTE_LEGAL }}  |  Nombre y firma'),
    (False, 'Representante de los trabajadores: {{ EMPRESA_REPRESENTANTE_TRABAJADORES }}  |  Nombre y firma'),
    (True, 'OTROS'),
    (False, 'Organización: {{ ORGANIZACION_NOMBRE }}'),
    (False, 'Fecha de emisión: {{ FECHA_EMISION }}  |  Vigencia: {{ VIGENCIA }}'),
    (False, 'LOGO: {{ LOGO }}'),
]


def build_template(path: Path) -> None:
    doc = Document()
    for is_heading, text in CONTENT:
        p = doc.add_paragraph()
        run = p.add_run(text)
        if is_heading:
            run.bold = True
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(11)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path.as_posix())


if __name__ == '__main__':
    out = Path('plantillas') / 'dc3_plantilla.docx'
    build_template(out)
    print(f'Wrote: {out.as_posix()}')


